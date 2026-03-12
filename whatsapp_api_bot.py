"""
WhatsApp Cloud API Bot — COD Label Generator
=============================================
This runs as a Flask webhook server. No browser needed.
Works with your phone's WhatsApp — send orders from anywhere.

SETUP (one-time):
1. Go to https://developers.facebook.com → Create App → Business → WhatsApp
2. In WhatsApp > API Setup:
   - Note your Phone Number ID, WhatsApp Business Account ID
   - Generate a permanent access token
3. Set up a webhook:
   - URL: https://your-server.com/webhook  (use ngrok for testing)
   - Verify token: pick any string, put it in VERIFY_TOKEN below
   - Subscribe to "messages" field
4. Fill in the config below and run:
   python whatsapp_api_bot.py

FREE TUNNEL (for testing without a server):
   ngrok http 5000
   Then paste the ngrok URL into Meta's webhook config.
"""

import re
import os
import sys
import json
import hashlib
import hmac
import time
import logging
import subprocess
import threading
import collections
import requests

from flask import Flask, request, jsonify

from docx import Document
from num2words import num2words
from copy import deepcopy

# Try importing docx2pdf (needs MS Word — only works on Windows with Word installed)
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None


# ============== CONFIG ==============
# Reads from environment variables first, falls back to hardcoded values.
# On Render, set these as Environment Variables in the dashboard.

WHATSAPP_TOKEN = os.environ.get("WHATSAPP_TOKEN", "")          # Your permanent access token
PHONE_NUMBER_ID = os.environ.get("PHONE_NUMBER_ID", "")        # Your WhatsApp Phone Number ID
VERIFY_TOKEN = os.environ.get("VERIFY_TOKEN", "cod_bot_verify")  # Webhook verification token
APP_SECRET = os.environ.get("APP_SECRET", "")                  # Optional: App secret for signature verification
ALLOWED_NUMBERS = []         # e.g. ["919342901848"] — leave empty to allow all
SEND_PROGRESS_REPLY = False  # False = don't send "Added X label(s)" after every order message

# ============== FILE CONFIG ==============

TEMPLATE = "cod_template.docx"
OUTPUT = "generated_labels.docx"
PROCESSED_FILE = "processed_orders_api.json"
BATCH_COUNTER_FILE = "batch_counter.txt"
STATE_FILE = "bot_state.json"
ORDERS_FILE = "collected_orders.json"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.environ.get("DATA_DIR", BASE_DIR)
os.makedirs(DATA_DIR, exist_ok=True)

TEMPLATE_PATH = os.environ.get("TEMPLATE_PATH", os.path.join(BASE_DIR, TEMPLATE))
OUTPUT_PATH = os.path.join(DATA_DIR, OUTPUT)
PROCESSED_PATH = os.path.join(DATA_DIR, PROCESSED_FILE)
BATCH_COUNTER_PATH = os.path.join(DATA_DIR, BATCH_COUNTER_FILE)
STATE_PATH = os.path.join(DATA_DIR, STATE_FILE)
ORDERS_PATH = os.path.join(DATA_DIR, ORDERS_FILE)

# ============== APP ==============

# --- Logging setup (visible in Render logs) ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    stream=sys.stdout,
)
logger = logging.getLogger(__name__)

app = Flask(__name__)


@app.route("/", methods=["GET"])
def health():
    """Health-check endpoint so Render knows the service is alive."""
    return jsonify({"status": "ok", "service": "cod-label-bot"}), 200


DELIM = r"[,:;.\s]"

# Message deduplication: track recently processed message IDs to ignore Meta retries.
# Uses an OrderedDict as a bounded cache (max 500 entries).
_processed_msg_ids_lock = threading.Lock()
_processed_msg_ids = collections.OrderedDict()  # msg_id -> timestamp
_MAX_MSG_CACHE = 500

# Maximum age (seconds) of a message we'll still process.
# Default 0 disables age-based dropping to avoid losing delayed but valid messages.
MAX_MESSAGE_AGE = int(os.environ.get("MAX_MESSAGE_AGE", "0"))

logger.info("=== COD Label Bot starting ===")
logger.info("WHATSAPP_TOKEN set: %s", bool(WHATSAPP_TOKEN))
logger.info("PHONE_NUMBER_ID set: %s", bool(PHONE_NUMBER_ID))
logger.info("APP_SECRET set: %s", bool(APP_SECRET))
logger.info("VERIFY_TOKEN: %s", VERIFY_TOKEN)
logger.info("TEMPLATE_PATH: %s (exists: %s)", TEMPLATE_PATH, os.path.exists(TEMPLATE_PATH))
logger.info("DATA_DIR: %s", DATA_DIR)
logger.info("=== Ready to receive webhooks ===")


# ---------------- STATE ----------------

def load_state():
    if os.path.exists(STATE_PATH):
        with open(STATE_PATH, "r") as f:
            return json.load(f)
    return {"collecting": False, "batch_count": 0}


def save_state(state):
    with open(STATE_PATH, "w") as f:
        json.dump(state, f)


def load_processed():
    if os.path.exists(PROCESSED_PATH):
        with open(PROCESSED_PATH, "r", encoding="utf-8") as f:
            return set(json.load(f))
    return set()


def save_processed(processed):
    with open(PROCESSED_PATH, "w", encoding="utf-8") as f:
        json.dump(list(processed), f)


def order_hash(data):
    key = f"{data['name']}|{data['phone']}|{data['pincode']}|{data['price']}|{data['item']}"
    return hashlib.sha256(key.encode()).hexdigest()[:16]


def load_orders():
    if os.path.exists(ORDERS_PATH):
        with open(ORDERS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_orders(orders):
    with open(ORDERS_PATH, "w", encoding="utf-8") as f:
        json.dump(orders, f, ensure_ascii=False)


def regenerate_docx(orders):
    """Rebuild the output DOCX from the current orders list."""
    if os.path.exists(OUTPUT_PATH):
        os.remove(OUTPUT_PATH)
    for data in orders:
        add_label(data)


def get_next_batch_number():
    num = 1
    if os.path.exists(BATCH_COUNTER_PATH):
        with open(BATCH_COUNTER_PATH, "r") as f:
            try:
                num = int(f.read().strip()) + 1
            except ValueError:
                num = 1
    with open(BATCH_COUNTER_PATH, "w") as f:
        f.write(str(num))
    return num


# ---------------- WHATSAPP API ----------------

def send_message(to, text):
    """Send a text message via WhatsApp Cloud API."""
    url = f"https://graph.facebook.com/v21.0/{PHONE_NUMBER_ID}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_TOKEN}",
        "Content-Type": "application/json",
    }
    payload = {
        "messaging_product": "whatsapp",
        "to": to,
        "type": "text",
        "text": {"body": text},
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=30)
    if resp.status_code != 200:
        print(f"  Send message failed: {resp.text}")
    return resp.status_code == 200


def send_document(to, file_path, caption=""):
    """Upload and send a document via WhatsApp Cloud API.
    Returns (success, error_message)."""
    # Step 1: Upload media
    url_upload = f"https://graph.facebook.com/v21.0/{PHONE_NUMBER_ID}/media"
    headers = {"Authorization": f"Bearer {WHATSAPP_TOKEN}"}

    filename = os.path.basename(file_path)
    mime = "application/pdf" if file_path.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    with open(file_path, "rb") as f:
        resp = requests.post(
            url_upload,
            headers=headers,
            files={"file": (filename, f, mime)},
            data={"messaging_product": "whatsapp"},
            timeout=60,
        )

    if resp.status_code != 200:
        err = f"Media upload failed: {resp.status_code} {resp.text}"
        print(f"  {err}")
        return False, err

    media_id = resp.json().get("id")
    if not media_id:
        err = f"Media upload response missing id: {resp.text}"
        print(f"  {err}")
        return False, err

    # Step 2: Send document message
    url_send = f"https://graph.facebook.com/v21.0/{PHONE_NUMBER_ID}/messages"
    payload = {
        "messaging_product": "whatsapp",
        "to": to,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": filename,
        },
    }
    if caption:
        payload["document"]["caption"] = caption

    resp = requests.post(
        url_send,
        headers={"Authorization": f"Bearer {WHATSAPP_TOKEN}", "Content-Type": "application/json"},
        json=payload,
        timeout=30,
    )
    if resp.status_code != 200:
        err = f"Send document failed: {resp.status_code} {resp.text}"
        print(f"  {err}")
        return False, err
    return True, ""


# ---------------- PARSING (same as whatsapp_label_bot.py) ----------------

def convert_words(price):
    words = num2words(price)
    words = words.replace(",", "")
    words = words.replace("-", " ")
    words = words.title()
    return words


def parse_order(text):
    text = text.strip()
    name_m = re.search(rf"Name{DELIM}+(.+)", text, re.I)
    pincode_m = re.search(rf"Pincode{DELIM}+\s*(\d{{6}})", text, re.I)
    state_m = re.search(rf"State{DELIM}+(.+)", text, re.I)
    phone_m = re.search(rf"Phone\s*(?:number)?{DELIM}+\s*(\d{{10}})", text, re.I)

    missing = []
    if not name_m:    missing.append("Name")
    if not pincode_m: missing.append("Pincode")
    if not state_m:   missing.append("State")
    if not phone_m:   missing.append("Phone")
    if missing:
        raise ValueError(f"Missing fields: {', '.join(missing)}")

    name = name_m.group(1).strip()
    pincode = pincode_m.group(1).strip()
    state = state_m.group(1).strip()
    phone = phone_m.group(1).strip()

    addr_m = re.search(rf"Address{DELIM}+(.+?)(?=\s*(?:City|Pincode))", text, re.S | re.I)
    address = addr_m.group(1).strip() if addr_m else ""
    address = re.sub(r"\s*\n\s*", ", ", address)
    address = address.strip(", ")
    address = re.sub(r",\s*,", ",", address)

    lines = text.strip().split("\n")
    price = None
    price_line_idx = None
    for idx in range(len(lines) - 1, -1, -1):
        stripped = lines[idx].strip()
        if re.match(r"^\d{2,6}$", stripped):
            price = int(stripped)
            price_line_idx = idx
            break

    if price is not None:
        phone_line_idx = None
        for idx, ln in enumerate(lines):
            if re.search(r"Phone", ln, re.I):
                phone_line_idx = idx
                break
        start = (phone_line_idx + 1) if phone_line_idx is not None else (price_line_idx - 1)
        while start < price_line_idx and not lines[start].strip():
            start += 1
        item_text = " ".join(lines[start:price_line_idx]).strip()
    else:
        item_price_m = re.search(r"(\d+\s*[A-Za-z].*?)[.;:]?\s+(\d{2,6})\s*$", text)
        if item_price_m:
            item_text = item_price_m.group(1).strip()
            price = int(item_price_m.group(2))
        else:
            raise ValueError("Could not parse item/price")

    item = parse_item_text(item_text)
    return {
        "name": name, "address": address, "state": state,
        "pincode": pincode, "phone": phone, "price": price, "item": item,
    }


def parse_item_text(item_text):
    cxe_m = re.match(r"(\d+)\s*([Cc][Xx][Ee])[.;:,]?\s*(.*)", item_text)
    if not cxe_m:
        code_m = re.match(r"(\d+)\s*(.*)", item_text)
        if code_m:
            qty = code_m.group(1)
            code = code_m.group(2).strip().upper().rstrip(".,;:")
            return f"{qty} {code}" if code else f"{qty} ITEM"
        return item_text.upper()

    cxe_qty = cxe_m.group(1)
    cxe_code = cxe_m.group(2).upper()
    rest = cxe_m.group(3).strip()
    items = [f"{cxe_qty} {cxe_code}"]
    if rest:
        parts = re.findall(r"(\d+\s*[A-Za-z][A-Za-z ]*)", rest)
        for p in parts:
            m = re.match(r"(\d+)\s*(.*)", p.strip())
            if m:
                items.append(f"{m.group(1)} {m.group(2).strip().title()}")
    return ", ".join(items)


def split_orders(text):
    parts = re.split(rf"(?=Name{DELIM})", text, flags=re.I)
    return [p.strip() for p in parts if p.strip() and re.search(rf"Name{DELIM}", p, re.I)]


# ---------------- LABEL GENERATION (same logic) ----------------

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def fill_table(table, data):
    price_words = convert_words(data["price"])
    placeholders = {
        "{{NAME}}": data["name"], "{{ADDRESS}}": data["address"],
        "{{STATE}}": data["state"], "{{PINCODE}}": data["pincode"],
        "{{PHONE}}": data["phone"], "{{PRICE}}": str(data["price"]),
        "{{PRICE_WORDS}}": price_words, "{{ITEM}}": data["item"],
    }
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text
                replaced = text
                for key, value in placeholders.items():
                    replaced = replaced.replace(key, value)
                if replaced != text:
                    if p.runs:
                        p.runs[0].text = replaced
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""
                    else:
                        p.text = replaced


def make_gap_element():
    gap = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "200")
    spacing.set(qn("w:after"), "200")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "12")
    rPr.append(sz)
    pPr.append(rPr)
    gap.append(pPr)
    return gap


def remove_empty_body_paragraphs(doc):
    body = doc.element.body
    for p in list(body.findall(qn("w:p"))):
        if p.getparent() != body:
            continue
        has_text = any(t.text for t in p.findall(f".//{qn('w:t')}") if t.text and t.text.strip())
        has_break = p.findall(f".//{qn('w:br')}")
        if not has_text and not has_break:
            body.remove(p)


def add_label(data):
    template_doc = Document(TEMPLATE_PATH)
    template_table = template_doc.tables[0]

    if os.path.exists(OUTPUT_PATH):
        doc = Document(OUTPUT_PATH)
    else:
        doc = Document(TEMPLATE_PATH)
        remove_empty_body_paragraphs(doc)
        fill_table(doc.tables[0], data)
        doc.save(OUTPUT_PATH)
        return

    tables_count = len(doc.tables)
    last_table_elem = doc.tables[-1]._element

    if tables_count % 2 == 0:
        page_break_p = OxmlElement("w:p")
        page_break_r = OxmlElement("w:r")
        page_break_br = OxmlElement("w:br")
        page_break_br.set(qn("w:type"), "page")
        page_break_r.append(page_break_br)
        page_break_p.append(page_break_r)
        last_table_elem.addnext(page_break_p)
        anchor = page_break_p
    else:
        anchor = last_table_elem

    gap = make_gap_element()
    anchor.addnext(gap)
    new_table = deepcopy(template_table._element)
    gap.addnext(new_table)
    fill_table(doc.tables[-1], data)
    doc.save(OUTPUT_PATH)


def stop_and_export():
    """Returns (pdf_path_or_None, docx_path)."""
    if not os.path.exists(OUTPUT_PATH):
        return None, None
    batch_num = get_next_batch_number()
    cod_docx = os.path.join(DATA_DIR, f"cod{batch_num}.docx")
    cod_pdf = os.path.join(DATA_DIR, f"cod{batch_num}.pdf")
    pdf_path = convert_to_pdf(OUTPUT_PATH)
    os.rename(OUTPUT_PATH, cod_docx)
    final_pdf = None
    if pdf_path and os.path.exists(pdf_path):
        os.rename(pdf_path, cod_pdf)
        final_pdf = cod_pdf
    for f in [PROCESSED_PATH, ORDERS_PATH]:
        if os.path.exists(f):
            os.remove(f)
    return final_pdf, cod_docx


def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    # Try docx2pdf (needs MS Word on Windows)
    if docx2pdf_convert:
        try:
            docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return pdf_path
        except Exception as e:
            print(f"  docx2pdf failed (needs MS Word): {e}")
    # Try LibreOffice (works on Linux servers / Windows if installed)
    try:
        outdir = os.path.dirname(os.path.abspath(docx_path)) or "."
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path
        print(f"  LibreOffice conversion failed: {result.stderr}")
    except FileNotFoundError:
        print("  LibreOffice not found on this machine")
    except Exception as e:
        print(f"  LibreOffice error: {e}")
    print("  PDF conversion unavailable — will send DOCX instead")
    return None


# ---------------- WEBHOOK ----------------

@app.route("/webhook", methods=["GET"])
def verify():
    """Meta sends a GET request to verify the webhook."""
    mode = request.args.get("hub.mode")
    token = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")
    logger.info("Webhook verify: mode=%s token_match=%s", mode, token == VERIFY_TOKEN)

    if mode == "subscribe" and token == VERIFY_TOKEN:
        logger.info("Webhook verified successfully.")
        return challenge, 200
    logger.warning("Webhook verification FAILED (token mismatch or wrong mode)")
    return "Forbidden", 403


def _is_duplicate_message(msg_id):
    """Return True if this message ID was already processed (dedup against Meta retries)."""
    with _processed_msg_ids_lock:
        if msg_id in _processed_msg_ids:
            return True
        _processed_msg_ids[msg_id] = time.time()
        # Evict oldest entries when cache is full
        while len(_processed_msg_ids) > _MAX_MSG_CACHE:
            _processed_msg_ids.popitem(last=False)
        return False


# Lock to serialize message processing (file I/O is not thread-safe)
_processing_lock = threading.Lock()


@app.route("/webhook", methods=["POST"])
def webhook():
    """Receive incoming WhatsApp messages.
    Returns 200 immediately, processes messages in a background thread
    so Meta doesn't time out and retry."""
    logger.info("Webhook POST received (%d bytes)", len(request.data or b""))

    # Verify signature if APP_SECRET is set
    if APP_SECRET:
        signature = request.headers.get("X-Hub-Signature-256", "")
        expected = "sha256=" + hmac.new(
            APP_SECRET.encode(), request.data, hashlib.sha256
        ).hexdigest()
        if not hmac.compare_digest(signature, expected):
            logger.warning("Invalid signature — rejecting request")
            return "Invalid signature", 403

    body = request.get_json()

    if not body:
        logger.info("Empty body — ignoring")
        return "OK", 200

    # Extract messages
    try:
        entry = body.get("entry", [{}])[0]
        changes = entry.get("changes", [{}])[0]
        value = changes.get("value", {})
        messages = value.get("messages", [])
    except (IndexError, KeyError):
        return "OK", 200

    # Collect valid messages to process
    now = int(time.time())
    to_process = []
    for msg in messages:
        if msg.get("type") != "text":
            continue

        msg_id = msg.get("id", "")
        sender = msg["from"]
        text = msg["text"]["body"].strip()
        msg_ts = int(msg.get("timestamp", "0"))

        # Skip duplicate messages (Meta retries)
        if msg_id and _is_duplicate_message(msg_id):
            print(f"  [SKIP] Duplicate message {msg_id} from {sender}")
            continue

        # Skip stale messages only when MAX_MESSAGE_AGE is configured (> 0)
        if MAX_MESSAGE_AGE > 0 and msg_ts and (now - msg_ts) > MAX_MESSAGE_AGE:
            print(f"  [SKIP] Stale message from {sender} ({now - msg_ts}s old): {text[:60]}")
            continue

        # Check allowed numbers
        if ALLOWED_NUMBERS and sender not in ALLOWED_NUMBERS:
            continue

        to_process.append((sender, text, msg_ts, msg_id))

    # Process in background thread so we return 200 immediately
    if to_process:
        thread = threading.Thread(target=_process_messages, args=(to_process,), daemon=True)
        thread.start()

    return "OK", 200


def _process_messages(messages_list):
    """Process messages in a background thread with a lock to serialize file access."""
    with _processing_lock:
        for sender, text, msg_ts, msg_id in messages_list:
            print(f"\n[MSG from {sender}]: {text[:100]}")
            try:
                handle_message(sender, text, msg_ts=msg_ts, msg_id=msg_id)
            except Exception as e:
                print(f"  Error handling message: {e}")
                try:
                    send_message(sender, f"Internal error while processing your message: {str(e)[:500]}")
                except Exception:
                    pass


def handle_message(sender, text, msg_ts=None, msg_id=None):
    """Process a single incoming message."""
    state = load_state()
    processed = load_processed()
    lower = text.strip().lower()
    if msg_ts is None:
        msg_ts = int(time.time())

    if not os.path.exists(TEMPLATE_PATH):
        send_message(
            sender,
            "Template file missing on server: cod_template.docx. "
            "Please deploy this file to Render and redeploy."
        )
        return

    # --- START command ---
    if lower == "start":
        state["collecting"] = True
        state["batch_count"] = 0
        state["collecting_started_at"] = msg_ts
        state["collecting_sender"] = sender
        # Clear previous batch data
        save_processed(set())
        save_orders([])
        if os.path.exists(OUTPUT_PATH):
            os.remove(OUTPUT_PATH)
        save_state(state)
        send_message(sender, "Started collecting orders.\nPaste order details now.\nSend 'stop' when done to get the PDF.\n\nCommands:\n  list — view all orders\n  delete <n> — remove order #n\n  stop — export PDF")
        print("  STARTED collecting")
        return

    # --- STOP command ---
    if lower == "stop":
        # Ignore delayed/replayed STOP older than the current collecting session.
        started_at = int(state.get("collecting_started_at", 0) or 0)
        if started_at and msg_ts and msg_ts < started_at:
            print(f"  [SKIP] Ignored stale STOP from {sender}")
            return

        # Only the sender who started the batch can stop it.
        collecting_sender = state.get("collecting_sender")
        if state.get("collecting") and collecting_sender and sender != collecting_sender:
            send_message(sender, "Only the number that sent 'start' can send 'stop' for this batch.")
            return

        if not state.get("collecting"):
            send_message(sender, "Not currently collecting. Send 'start' first.")
            return

        count = state.get("batch_count", 0)
        state["collecting"] = False
        state["batch_count"] = 0
        save_state(state)

        if count == 0:
            send_message(sender, "No orders were collected. Nothing to export.")
            return

        send_message(sender, f"Stopped. {count} label(s) collected.\nGenerating PDF...")

        pdf_path, docx_path = stop_and_export()
        if pdf_path:
            batch_num = int(open(BATCH_COUNTER_PATH).read().strip())
            success, send_err = send_document(sender, pdf_path, caption=f"COD Labels — Batch {batch_num}")
            if success:
                send_message(sender, f"PDF sent! (cod{batch_num}.pdf)\nSend 'start' for the next batch.")
            else:
                send_message(sender, f"PDF saved locally but could not send.\nReason: {send_err[:700]}")
        elif docx_path:
            batch_num = int(open(BATCH_COUNTER_PATH).read().strip())
            success, send_err = send_document(sender, docx_path, caption=f"COD Labels — Batch {batch_num}")
            if success:
                send_message(sender, f"DOCX file sent! (cod{batch_num}.docx)\nInstall MS Word or LibreOffice for PDF.\nSend 'start' for the next batch.")
            else:
                send_message(sender, f"Could not send file.\nReason: {send_err[:700]}")
        else:
            send_message(sender, "No labels found to export.")

        print(f"  STOPPED — {count} labels exported")
        return

    # --- LIST command ---
    if lower == "list":
        if not state.get("collecting"):
            send_message(sender, "Not currently collecting. Send 'start' first.")
            return
        collecting_sender = state.get("collecting_sender")
        if collecting_sender and sender != collecting_sender:
            send_message(sender, "This batch belongs to another number.")
            return
        orders = load_orders()
        if not orders:
            send_message(sender, "No orders collected yet.")
            return
        lines = [f"Orders collected: {len(orders)}\n"]
        for i, o in enumerate(orders, 1):
            lines.append(f"{i}. {o['name']} — {o['phone']} — ₹{o['price']} — {o['item']}")
        lines.append("\nSend 'delete <n>' to remove an order.")
        send_message(sender, "\n".join(lines))
        return

    # --- DELETE command ---
    delete_m = re.match(r"delete\s+(\d+)", lower)
    if delete_m:
        if not state.get("collecting"):
            send_message(sender, "Not currently collecting. Send 'start' first.")
            return
        collecting_sender = state.get("collecting_sender")
        if collecting_sender and sender != collecting_sender:
            send_message(sender, "This batch belongs to another number.")
            return
        idx = int(delete_m.group(1))
        orders = load_orders()
        if idx < 1 or idx > len(orders):
            send_message(sender, f"Invalid order number. You have {len(orders)} order(s). Send 'list' to see them.")
            return
        removed = orders.pop(idx - 1)
        save_orders(orders)
        # Rebuild processed hashes and DOCX
        save_processed({order_hash(o) for o in orders})
        regenerate_docx(orders)
        state["batch_count"] = len(orders)
        save_state(state)
        send_message(sender, f"Deleted order #{idx}: {removed['name']} — ₹{removed['price']}\n{len(orders)} order(s) remaining.")
        print(f"  DELETED order #{idx}: {removed['name']}")
        return

    # --- STATUS command ---
    if lower == "status":
        collecting = state.get("collecting", False)
        count = state.get("batch_count", 0)
        if collecting:
            owner = state.get("collecting_sender", "")
            if owner and sender != owner:
                send_message(sender, "A batch is currently active for another number.")
            else:
                send_message(sender, f"Collecting orders: {count} label(s) so far.\nSend 'stop' to export PDF.")
        else:
            send_message(sender, "Not collecting. Send 'start' to begin.")
        return

    # --- Order data ---
    if not state.get("collecting"):
        # Ignore replayed order blocks when not collecting to avoid repeated spam replies.
        if re.search(rf"Name{DELIM}", text, re.I):
            print(f"  [SKIP] Order text while not collecting (likely replay): {text[:60]}")
            return
        send_message(sender, "Send 'start' first to begin collecting orders.")
        return

    collecting_sender = state.get("collecting_sender")
    if collecting_sender and sender != collecting_sender:
        send_message(sender, "This batch belongs to another number. Send 'start' to begin your own batch.")
        return

    # Try to parse orders from the message
    if not re.search(rf"Name{DELIM}", text, re.I):
        return  # Not an order, ignore silently

    order_blocks = split_orders(text)
    added = 0

    orders = load_orders()

    for block in order_blocks:
        try:
            data = parse_order(block)
            h = order_hash(data)
            if h in processed:
                continue
            add_label(data)
            processed.add(h)
            orders.append(data)
            state["batch_count"] = state.get("batch_count", 0) + 1
            added += 1
            print(f"  + Label for: {data['name']} (₹{data['price']})")
        except ValueError as e:
            send_message(sender, f"Could not parse order: {e}")

    save_orders(orders)
    save_processed(processed)
    save_state(state)

    if added and SEND_PROGRESS_REPLY:
        total = state["batch_count"]
        send_message(sender, f"Added {added} label(s). Total this batch: {total}\nSend more orders or 'stop' to export.")


# ---------------- MAIN ----------------

if __name__ == "__main__":
    if not WHATSAPP_TOKEN or not PHONE_NUMBER_ID:
        print("=" * 55)
        print("  WhatsApp Cloud API Bot — COD Label Generator")
        print("=" * 55)
        print()
        print("ERROR: Please fill in your API credentials in the")
        print("config section at the top of this file:")
        print("  WHATSAPP_TOKEN   = your access token")
        print("  PHONE_NUMBER_ID  = your phone number ID")
        print()
        print("See setup instructions at the top of this file.")
        sys.exit(1)

    print("=" * 55)
    print("  WhatsApp Cloud API Bot — COD Label Generator")
    print("=" * 55)
    print()
    print("Commands (send via WhatsApp):")
    print("  start    — Begin collecting orders")
    print("  stop     — Export PDF & send it back")
    print("  status   — Check current batch count")
    print()
    print("Starting webhook server on port 5000...")
    print(f"Data directory: {DATA_DIR}")
    print("If ngrok shows ERR_NGROK_4018, run once:")
    print("  .\\ngrok.exe config add-authtoken <YOUR_NGROK_TOKEN>")
    print("Then start tunnel:")
    print("  .\\ngrok.exe http 5000")
    print()

    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
